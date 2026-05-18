"""
Generador de documentos Word (.docx) para el informe municipal.

Produce un archivo formal listo para imprimir con:
- Encabezado con logos institucionales
- Tablas de distribución y estado de lotes
- Gráficos matplotlib embebidos
- Análisis de observaciones
- (Opcional) Detalle por manzana
"""
from __future__ import annotations

import io
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from src.report.charts import (
    grafico_barras_por_manzana_matplotlib,
    grafico_torta_estados_matplotlib,
)
from src.report.narrative import armar_narrativa

# Ruta base del proyecto (dos niveles arriba de este archivo: src/report → raíz)
_BASE_DIR = Path(__file__).resolve().parent.parent.parent
_LOGO_MUN = _BASE_DIR / "assets" / "logos" / "placeholder_municipalidad.png"
_LOGO_SPVH = _BASE_DIR / "assets" / "logos" / "placeholder_spvh.png"

# ─── Helpers de estilo ────────────────────────────────────────────────────────

def _set_font(run, size_pt: int = 11, bold: bool = False, color_hex: str | None = None) -> None:
    """Aplica fuente Calibri con tamaño y estilo al run."""
    run.font.name = "Calibri"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_hex:
        r, g, b = int(color_hex[0:2], 16), int(color_hex[2:4], 16), int(color_hex[4:6], 16)
        run.font.color.rgb = RGBColor(r, g, b)


def _paragraph_font(paragraph, size_pt: int = 11) -> None:
    """Ajusta fuente y tamaño base del párrafo (afecta a todos sus runs)."""
    for run in paragraph.runs:
        _set_font(run, size_pt)


def _set_cell_bg(cell, hex_color: str) -> None:
    """Pinta el fondo de una celda de tabla con color hex (sin #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _add_horizontal_rule(doc: Document) -> None:
    """Agrega un párrafo con borde inferior (línea horizontal)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4A4A4A")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def _set_page_margins(doc: Document, margin_cm: float = 2.5) -> None:
    """Establece márgenes uniformes en todas las secciones."""
    for section in doc.sections:
        section.top_margin = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin = Cm(margin_cm)
        section.right_margin = Cm(margin_cm)


def _add_section_heading(doc: Document, text: str) -> None:
    """Agrega un encabezado de sección en negrita 12pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    _set_font(run, size_pt=12, bold=True)


# ─── Construcción del documento ───────────────────────────────────────────────

def _build_header(doc: Document) -> None:
    """Tabla de encabezado con logos institucionales (2 columnas)."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    left_cell, right_cell = table.columns[0].cells[0], table.columns[1].cells[0]

    # Ajustar ancho de columnas (~6 cm cada una)
    for cell in (left_cell, right_cell):
        cell.width = Cm(6)

    # Insertar logos si los archivos existen
    for cell, logo_path in ((left_cell, _LOGO_MUN), (right_cell, _LOGO_SPVH)):
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_path.exists():
            run = p.add_run()
            run.add_picture(str(logo_path), width=Cm(5))
        else:
            run = p.add_run(logo_path.stem)
            _set_font(run, size_pt=10)

    _add_horizontal_rule(doc)


def _build_title_date(doc: Document, titulo: str, fecha: str) -> None:
    """Título centrado en negrita y fecha alineada a la derecha."""
    # Título
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(4)
    run = p_title.add_run(titulo)
    _set_font(run, size_pt=14, bold=True)

    # Fecha
    p_fecha = doc.add_paragraph()
    p_fecha.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_fecha.paragraph_format.space_after = Pt(8)
    run = p_fecha.add_run(f"Fecha: {fecha}")
    _set_font(run, size_pt=11)


def _build_introduccion(doc: Document, texto: str) -> None:
    """Párrafo de introducción justificado."""
    p = doc.add_paragraph(texto)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    _paragraph_font(p, size_pt=11)


def _build_distribucion_manzanas(doc: Document, por_manzana: pd.DataFrame, total_lotes: int) -> None:
    """Tabla de distribución de lotes por manzana con fila de total."""
    _add_section_heading(doc, "Distribución de Lotes")

    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"

    # Encabezado
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Manzana", "Total Lotes")):
        p = cell.paragraphs[0]
        run = p.add_run(text)
        _set_font(run, size_pt=11, bold=True)
        _set_cell_bg(cell, "D1D5DB")

    # Filas de datos
    for _, row in por_manzana.iterrows():
        cells = table.add_row().cells
        run0 = cells[0].paragraphs[0].add_run(str(row["MANZANA"]))
        run1 = cells[1].paragraphs[0].add_run(str(row["TOTAL"]))
        _set_font(run0, size_pt=11)
        _set_font(run1, size_pt=11)

    # Fila total
    total_cells = table.add_row().cells
    run_t = total_cells[0].paragraphs[0].add_run("Total")
    run_v = total_cells[1].paragraphs[0].add_run(str(total_lotes))
    _set_font(run_t, size_pt=11, bold=True)
    _set_font(run_v, size_pt=11, bold=True)
    for c in total_cells:
        _set_cell_bg(c, "E5E7EB")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_estado_general(doc: Document, estado_general: dict) -> None:
    """Tabla de estado general con porcentajes; fila Escriturados en azul claro."""
    _add_section_heading(doc, "Estado General")

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # Encabezado
    hdr = table.rows[0].cells
    for cell, text in zip(hdr, ("Estado", "Cantidad", "%")):
        run = cell.paragraphs[0].add_run(text)
        _set_font(run, size_pt=11, bold=True)
        _set_cell_bg(cell, "D1D5DB")

    for grupo, datos in estado_general.items():
        cells = table.add_row().cells
        run_g = cells[0].paragraphs[0].add_run(grupo)
        run_c = cells[1].paragraphs[0].add_run(str(datos["cantidad"]))
        run_p = cells[2].paragraphs[0].add_run(f"{datos['porcentaje']}%")
        for run in (run_g, run_c, run_p):
            _set_font(run, size_pt=11)
        # Resaltar Escriturados en azul claro
        if grupo == "Escriturados":
            for c in cells:
                _set_cell_bg(c, "DBEAFE")

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_charts(doc: Document, estado_general: dict, estado_por_manzana_df: pd.DataFrame, nombre_barrio: str) -> None:
    """Inserta gráfico de torta y de barras como imágenes PNG."""
    # Pie chart
    pie_buf = grafico_torta_estados_matplotlib(
        estado_general,
        titulo=f"Estado General de Lotes – Barrio {nombre_barrio}",
    )
    p_pie = doc.add_paragraph()
    p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_pie = p_pie.add_run()
    run_pie.add_picture(pie_buf, width=Cm(10))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Bar chart
    bar_buf = grafico_barras_por_manzana_matplotlib(
        estado_por_manzana_df,
        titulo=f"Estado de Lotes por Manzana – Barrio {nombre_barrio}",
    )
    p_bar = doc.add_paragraph()
    p_bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_bar = p_bar.add_run()
    run_bar.add_picture(bar_buf, width=Cm(14))

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_observaciones(doc: Document, patrones_obs: dict) -> None:
    """Lista con viñetas de patrones detectados en observaciones."""
    if not patrones_obs:
        return

    _add_section_heading(doc, "Análisis de Observaciones")

    for patron, cantidad in patrones_obs.items():
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{patron}: {cantidad} caso{'s' if cantidad != 1 else ''}")
        _set_font(run, size_pt=11)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _build_detalle_manzanas(doc: Document, df: pd.DataFrame) -> None:
    """Detalle completo por manzana: encabezado + tabla de lotes."""
    _add_section_heading(doc, "Detalle por Manzana")

    manzanas = sorted(df["MANZANA"].unique())
    columnas = ["TITULAR", "DIRECCION", "ESTADO", "OBSERVACIONES"]

    for manzana in manzanas:
        # Sub-encabezado de manzana
        p_mz = doc.add_paragraph()
        p_mz.paragraph_format.space_before = Pt(8)
        p_mz.paragraph_format.space_after = Pt(2)
        run_mz = p_mz.add_run(str(manzana))
        _set_font(run_mz, size_pt=11, bold=True)

        subset = df[df["MANZANA"] == manzana]

        table = doc.add_table(rows=1, cols=len(columnas))
        table.style = "Table Grid"

        # Fila de encabezado
        hdr_cells = table.rows[0].cells
        for cell, col in zip(hdr_cells, columnas):
            run = cell.paragraphs[0].add_run(col)
            _set_font(run, size_pt=9, bold=True)
            _set_cell_bg(cell, "D1D5DB")

        # Filas de datos
        for _, row in subset.iterrows():
            data_cells = table.add_row().cells
            for cell, col in zip(data_cells, columnas):
                valor = row[col]
                text = str(valor) if valor is not None and not (isinstance(valor, float) and pd.isna(valor)) else ""
                run = cell.paragraphs[0].add_run(text)
                _set_font(run, size_pt=9)

        doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ─── Función pública ──────────────────────────────────────────────────────────

def generar_word(
    df: pd.DataFrame,
    nombre_barrio: str,
    fecha: str,
    incluir_detalle_manzanas: bool = False,
) -> io.BytesIO:
    """
    Genera un documento Word (.docx) formal con el informe del barrio.

    Parámetros
    ----------
    df:
        DataFrame unificado (salida de ``cargar_excel``).
    nombre_barrio:
        Nombre del barrio para encabezados y textos.
    fecha:
        Fecha del informe en formato libre.
    incluir_detalle_manzanas:
        Si True, agrega una sección con el detalle completo por manzana.

    Retorna
    -------
    io.BytesIO
        Buffer con el archivo .docx listo para descargar o guardar.
    """
    narrativa = armar_narrativa(df, nombre_barrio, fecha)

    doc = Document()
    _set_page_margins(doc, margin_cm=2.5)

    # Fuente base del documento
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    # 1. Encabezado con logos
    _build_header(doc)

    # 2. Título y fecha
    _build_title_date(doc, narrativa["titulo"], narrativa["fecha"])

    # 3. Introducción
    _build_introduccion(doc, narrativa["introduccion"])

    # 4. Distribución de lotes por manzana
    _build_distribucion_manzanas(doc, narrativa["por_manzana"], narrativa["total_lotes"])

    # 5. Estado general
    _build_estado_general(doc, narrativa["estado_general"])

    # 6. Gráficos
    _build_charts(doc, narrativa["estado_general"], narrativa["estado_por_manzana"], nombre_barrio)

    # 7. Análisis de observaciones
    _build_observaciones(doc, narrativa["patrones_obs"])

    # 8. (Opcional) Detalle por manzana
    if incluir_detalle_manzanas:
        _build_detalle_manzanas(doc, df)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
